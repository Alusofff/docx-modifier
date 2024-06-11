from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
def change_font_and_size(document, font_name='Times New Roman', font_size=14):
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # For the font to be applied to non-English text
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), font_name)
def set_line_spacing(document, line_spacing=1.5):
    for paragraph in document.paragraphs:
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(int(line_spacing * 240)))  # 1.5 line spacing
        pPr.append(spacing)
def modify_docx(file_path, font_name='Times New Roman', font_size=14, line_spacing=1.5):
    document = Document(file_path)
    change_font_and_size(document, font_name, font_size)
    set_line_spacing(document, line_spacing)
    document.save(file_path)
